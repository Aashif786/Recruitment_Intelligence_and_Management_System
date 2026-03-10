from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, BackgroundTasks, Form
from sqlalchemy.orm import Session, joinedload
import os
import json
from datetime import datetime, timezone
from app.database import get_db
from app.models import User, Application, Job, ResumeExtraction, Interview, InterviewAnswer
from app.schemas import ApplicationCreate, ApplicationStatusUpdate, ApplicationResponse, ApplicationDetailResponse
from app.auth import get_current_user, get_current_hr
from app.services.ai_service import parse_resume_with_ai
from app.services.email_service import send_application_received_email, send_rejected_email, send_approved_for_interview_email
import secrets
from passlib.context import CryptContext
from datetime import datetime, timedelta

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

from app.config import get_settings
settings = get_settings()

UPLOAD_DIR = settings.uploads_dir / "resumes"
PHOTO_DIR = settings.uploads_dir / "photos"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
PHOTO_DIR.mkdir(parents=True, exist_ok=True)

router = APIRouter(prefix="/api/applications", tags=["applications"])

<<<<<<< HEAD
=======
@router.get("/ranking/{job_id}")
def get_candidate_ranking(
    job_id: int,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get ranked candidates for a specific job (Point 3)"""
    from app.services.candidate_service import CandidateService
    service = CandidateService(db)
    ranked = service.get_ranked_candidates(job_id)
    
    result = []
    for idx, app in enumerate(ranked):
        result.append({
            "rank": idx + 1,
            "id": app.id,
            "candidate_name": app.candidate_name,
            "composite_score": app.composite_score,
            "recommendation": app.recommendation,
            "status": app.status
        })
    return result


>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
@router.post("/apply", response_model=ApplicationResponse)
async def apply_for_job(
    job_id: int = Form(...),
    candidate_name: str = Form(...),
    candidate_email: str = Form(...),
    candidate_phone: str = Form(None),
    resume_file: UploadFile = File(...),
    photo_file: UploadFile = File(None),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db)
):
    """Apply for a job with resume (Public endpoint)"""
    # Check if job exists and is open
    job = db.query(Job).filter(Job.id == job_id, Job.status == "open").first()
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found or not open"
        )
    
    # Check if already applied
    # Check if already applied
    candidate_email = candidate_email.lower().strip()
    existing_app = db.query(Application).filter(
        Application.job_id == job_id,
        Application.candidate_email == candidate_email
    ).first()
    
    if existing_app:
        # If the previous application was rejected, allow re-application by deleting the old one
        if existing_app.status == "rejected":
            try:
                # Start fresh - delete the old application tree (cascades should handle relations)
                db.delete(existing_app)
                db.commit()
            except Exception as e:
                db.rollback()
                raise HTTPException(status_code=500, detail="Failed to recycle rejected application securely")
            # Loop continues to create new app
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="You have already applied for this job"
            )
    
    MAX_FILE_SIZE = 5 * 1024 * 1024 # 5MB
    ALLOWED_TYPES = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/plain"]
    
    # Validate file content type
    if resume_file.content_type not in ALLOWED_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid file type. Only PDF and DOCX allowed."
        )
        
    # Validate file size (Need to read chunk to be safe, but spooled file has .size or we check after reading)
    # Since UploadFile is spooled, we can check size if headers provided, or read content.
    content = await resume_file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File too large. Maximum size is 5MB."
        )
            
    # Save resume file
    file_extension = resume_file.filename.split(".")[-1]
    safe_email = candidate_email.replace('@', '_').replace('.', '_')
    filename = f"{safe_email}_{job_id}_{datetime.now(timezone.utc).timestamp()}.{file_extension}"
<<<<<<< HEAD
    file_path = os.path.join(UPLOAD_DIR, filename).replace("\\", "/")
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Save photo file if provided
    photo_path = None
=======
    
    # Absolute path for saving the file
    abs_file_path = os.path.join(UPLOAD_DIR, filename).replace("\\", "/")
    # Relative path for storing in DB (starts with 'uploads/')
    rel_file_path = f"uploads/resumes/{filename}"
    
    with open(abs_file_path, "wb") as f:
        f.write(content)
    
    # Save photo file if provided
    rel_photo_path = None
>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
    if photo_file:
        photo_content = await photo_file.read()
        if len(photo_content) > MAX_FILE_SIZE:
             raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Photo too large. Maximum size is 5MB."
            )
        
        photo_ext = photo_file.filename.split(".")[-1]
        photo_filename = f"photo_{safe_email}_{job_id}_{datetime.now(timezone.utc).timestamp()}.{photo_ext}"
<<<<<<< HEAD
        photo_path = os.path.join(PHOTO_DIR, photo_filename).replace("\\", "/")
        with open(photo_path, "wb") as f:
=======
        
        # Absolute path for saving
        abs_photo_path = os.path.join(PHOTO_DIR, photo_filename).replace("\\", "/")
        # Relative path for DB
        rel_photo_path = f"uploads/photos/{photo_filename}"
        
        with open(abs_photo_path, "wb") as f:
>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
            f.write(photo_content)
    
    # Create application
    new_application = Application(
        job_id=job_id,
        candidate_name=candidate_name,
        candidate_email=candidate_email,
        candidate_phone=candidate_phone,
<<<<<<< HEAD
        resume_file_path=file_path,
        resume_file_name=resume_file.filename,
        candidate_photo_path=photo_path,
=======
        resume_file_path=rel_file_path,
        resume_file_name=resume_file.filename,
        candidate_photo_path=rel_photo_path,
>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
        status="submitted"
    )
    
    try:
        db.add(new_application)
        db.commit()
        db.refresh(new_application)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to save new application securely")
    
<<<<<<< HEAD
    # Parse resume with AI (async in background would be better)
    try:
        # Read resume file
        # Parse resume text based on file type
        try:
            resume_text = ""
            file_ext = file_path.lower().split('.')[-1]
            
            if file_ext == 'pdf':
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    for page in reader.pages:
                        resume_text += page.extract_text() + "\n"
                except Exception as e:
                    print(f"PDF Error: {e}")
                    # Fallback to binary decode if PDF read fails (unlikely to work but last resort)
                    with open(file_path, "rb") as f:
                        resume_text = f.read().decode('utf-8', errors='ignore')
                        
            elif file_ext in ['docx', 'doc']:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    for para in doc.paragraphs:
                        resume_text += para.text + "\n"
                except Exception as e:
                    print(f"DOCX Error: {e}")
                    with open(file_path, "rb") as f:
                        resume_text = f.read().decode('utf-8', errors='ignore')
                        
            else:
                # Text file
                with open(file_path, "rb") as f:
                    resume_text = f.read().decode('utf-8', errors='ignore')
                    
            if not resume_text.strip():
                resume_text = "No readable text found in resume."
                
        except Exception as e:
            print(f"Text Extraction Error: {e}")
            resume_text = "Error extracting text."
        
        # Parse with AI
        extraction_data = await parse_resume_with_ai(
            resume_text,
            job.id,
            job.description
        )
        
        # Store extraction
        resume_extraction = ResumeExtraction(
            application_id=new_application.id,
            extracted_text=resume_text,  # Store FULL extracted text
            summary=extraction_data.get("summary", ""), # Store AI summary
=======
    # Move all heavy processing to background task to prevent timeouts (Point 10 - Robustness)
    background_tasks.add_task(
        process_application_background, 
        new_application.id, 
        job_id, 
        abs_file_path, 
        candidate_email, 
        candidate_name
    )
    
    return new_application

async def process_application_background(application_id: int, job_id: int, abs_file_path: str, candidate_email: str, candidate_name: str):
    """Heavy AI processing and notification workflow in background"""
    db = SessionLocal()
    try:
        from app.services.candidate_service import CandidateService
        cand_service = CandidateService(db)
        
        # Reload objects in this session
        application = db.query(Application).filter(Application.id == application_id).first()
        job = db.query(Job).filter(Job.id == job_id).first()
        if not application or not job:
            db.close()
            return

        # 1. Initial State
        cand_service.advance_stage(application_id, "Application Submitted", "pass")
        cand_service.create_audit_log(None, "APPLICATION_SUBMITTED", "Application", application_id, {"email": candidate_email})
        
        # 2. Screening Stage
        cand_service.advance_stage(application_id, "Resume Screening", "pending")
        
        # Parse resume text based on file type
        resume_text = ""
        try:
            file_ext = abs_file_path.lower().split('.')[-1]
            if file_ext == 'pdf':
                from pypdf import PdfReader
                reader = PdfReader(abs_file_path)
                for page in reader.pages:
                    resume_text += page.extract_text() + "\n"
            elif file_ext in ['docx', 'doc']:
                import docx
                doc = docx.Document(abs_file_path)
                for para in doc.paragraphs:
                    resume_text += para.text + "\n"
            else:
                with open(abs_file_path, "rb") as f:
                    resume_text = f.read().decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"Background Text Extraction Error: {e}")
            cand_service.create_audit_log(None, "RESUME_TEXT_EXTRACTION_FAILED", "Application", application_id, {"error": str(e)})
            resume_text = "Error extracting text."
        
        if not resume_text.strip():
            resume_text = "No readable text found."

        # AI Parsing
        extraction_data = await parse_resume_with_ai(resume_text, job_id, job.description)
        
        # Store extraction
        resume_extraction = ResumeExtraction(
            application_id=application_id,
            extracted_text=resume_text,
            summary=extraction_data.get("summary", ""),
>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
            extracted_skills=json.dumps(extraction_data.get("skills") or []),
            years_of_experience=extraction_data.get("experience"),
            education=json.dumps(extraction_data.get("education") or []),
            previous_roles=json.dumps(extraction_data.get("roles") or []),
            experience_level=extraction_data.get("experience_level"),
            resume_score=extraction_data.get("score", 0),
            skill_match_percentage=extraction_data.get("match_percentage", 0)
        )
        db.add(resume_extraction)
        
<<<<<<< HEAD
        # --- Validation Logic ---
        rejection_reasons = []
        
        # 0. Check if it's a resume
        if extraction_data.get("is_resume") is False:
             rejection_reasons.append("uploaded document is not a resume")
        
        # 1. Check for parsing failure
        # Heuristic: If skills are empty or extracted text indicates failure
        if not extraction_data.get("skills") and extraction_data.get("experience") == 0:
             rejection_reasons.append("resume parsing failed")
             
        # 2. Check for experience level mismatch
        # Normalize levels for comparison
        job_level = job.experience_level.lower().strip() if job.experience_level else ""
        candidate_level = str(extraction_data.get("experience_level", "")).lower().strip()
        
        # Define hierarchy
        levels = {
            "intern": 0,
            "junior": 1,
            "mid": 2, "mid-level": 2,
            "senior": 3,
            "lead": 4, "manager": 4, "lead / manager": 4
        }
        
        job_level_rank = levels.get(job_level, -1)
        candidate_level_rank = levels.get(candidate_level, -1)
        
        # If both levels are recognized, check if candidate is lower than required
        if job_level_rank != -1 and candidate_level_rank != -1:
            if candidate_level_rank < job_level_rank:
                 rejection_reasons.append("experience level mismatch")
        
        # If rejected, update status
        raw_access_key = None
        if rejection_reasons:
            new_application.status = "rejected"
            new_application.hr_notes = f"Auto-rejected based on: {', '.join(rejection_reasons)}"
        else:
            new_application.status = "approved_for_interview"
            new_application.hr_notes = "Auto-approved for interview."
            
            import uuid
            raw_access_key = secrets.token_urlsafe(16)
            hashed_key = pwd_context.hash(raw_access_key)
            expiration = datetime.now(timezone.utc) + timedelta(hours=24)
            unique_test_id = f"TEST-{uuid.uuid4().hex[:8].upper()}"
            
            new_interview = Interview(
                test_id=unique_test_id,
                application_id=new_application.id,
                status='not_started',
                access_key_hash=hashed_key,
                expires_at=expiration,
                is_used=False
            )
            db.add(new_interview)
            
        try:
            db.commit()
        except Exception as e:
            db.rollback()
            print(f"Error committing auto-rejection or approval: {e}")

        # Send notification to HR
        from app.models import Notification
        try:
            notification_type = "new_application"
            message = f"{candidate_name} has applied for the {job.title} position."
            
            if new_application.status == "rejected":
                message += f" (Auto-rejected: {', '.join(rejection_reasons)})"
                background_tasks.add_task(send_rejected_email, candidate_email, job.title, True)
            else:
                message += " (Auto-approved for interview)"
                # Generate Interview Access Key
                raw_access_key = secrets.token_urlsafe(16)
                hashed_key = pwd_context.hash(raw_access_key)
                expiration = datetime.now(timezone.utc)() + timedelta(hours=24)
                
                new_interview = Interview(
                    application_id=new_application.id,
                    status='not_started',
                    access_key_hash=hashed_key,
                    expires_at=expiration,
                    is_used=False
                )
                db.add(new_interview)
                db.commit()
                
                background_tasks.add_task(send_application_received_email, candidate_email, job.title)
                background_tasks.add_task(send_approved_for_interview_email, candidate_email, job.title, raw_access_key)
            
            notification = Notification(
                user_id=job.hr_id,
                notification_type=notification_type,
                title=f"New Application: {candidate_name}",
                message=message,
                related_application_id=new_application.id
            )
            db.add(notification)
            db.commit()
        except Exception as e:
            db.rollback()
            print(f"Error creating notification: {e}")

    except Exception as e:
        print(f"Error parsing resume: {e}")
        # Application is still created, just resume parsing failed
        background_tasks.add_task(send_application_received_email, candidate_email, job.title)
    
    return new_application

@router.get("", response_model=list[ApplicationDetailResponse])
def get_hr_applications(
=======
        # Update Application summary fields
        application.resume_score = extraction_data.get("score", 0)
        db.commit()

        # Recommendation and Progression
        res_score_norm = extraction_data.get("score", 0) * 10
        match_perc = extraction_data.get("match_percentage", 0)
        avg_score = (res_score_norm + match_perc) / 2
        
        status = "pass" if avg_score >= 70 else ("fail" if avg_score < 40 else "hold")
        note = "Strong match - automated progression" if status == "pass" else ("Low compatibility" if status == "fail" else "Manual review required")
        
        cand_service.advance_stage(application_id, "Resume Screening", status, avg_score, note)
        cand_service.create_audit_log(None, "RESUME_SCREENING_COMPLETED", "Application", application_id, {"avg_score": avg_score, "status": status})

        # Notifications
        if status == "pass":
            # Direct pass to aptitude
            raw_access_key = secrets.token_urlsafe(16)
            hashed_key = pwd_context.hash(raw_access_key) # Use pwd_context.hash
            expiration = datetime.now(timezone.utc) + timedelta(hours=24)
            
            new_interview = Interview(
                test_id=f"TEST-{secrets.token_hex(4).upper()}",
                application_id=application_id,
                status='not_started',
                access_key_hash=hashed_key,
                expires_at=expiration,
                is_used=False # Added missing field
            )
            db.add(new_interview)
            application.status = "approved_for_interview" # Update application status
            await send_approved_for_interview_email(candidate_email, job.title, raw_access_key)
        elif status == "fail":
            application.status = "rejected" # Update application status
            await send_rejected_email(candidate_email, job.title, True) # Auto-rejected
        else: # hold
            application.status = "submitted" # Keep as submitted for manual review
            await send_application_received_email(candidate_email, job.title)
            
        db.commit()
    except Exception as e:
        print(f"CRITICAL Background Error processing application {application_id}: {e}")
        db.rollback()
        # Log the critical error
        try:
            cand_service = CandidateService(db) # Re-initialize if needed, or pass db
            cand_service.create_audit_log(None, "BACKGROUND_PROCESSING_FAILED", "Application", application_id, {"error": str(e)})
            # Optionally update application status to indicate processing failed
            application = db.query(Application).filter(Application.id == application_id).first()
            if application:
                application.status = "processing_failed"
                application.hr_notes = f"Automated processing failed: {e}"
                db.commit()
        except Exception as log_e:
            print(f"Failed to log critical error for application {application_id}: {log_e}")
    finally:
        db.close()
    
    # The return value of a background task is not used by FastAPI.
    # The original snippet returned new_application, but it's not necessary here.
    # Keeping it for consistency with the provided snippet, but it has no effect.
    return application 

@router.get("", response_model=list[ApplicationDetailResponse])
def get_hr_applications(
    job_id: int = None,
>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get all applications for HR's jobs (HR only)"""
<<<<<<< HEAD
    applications = db.query(Application).join(Job).options(
        joinedload(Application.job),
        joinedload(Application.resume_extraction),
        joinedload(Application.interview)
    ).all()
=======
    # Join with stages for visualization (Point 12)
    # Use outerjoin to ensure apps with missing jobs (shouldn't happen but safe) still show or at least don't crash
    query = db.query(Application).outerjoin(Job).options(
        joinedload(Application.job),
        joinedload(Application.resume_extraction),
        joinedload(Application.interview),
        joinedload(Application.pipeline_stages)
    )
    
    # Filter by job if requested
    if job_id:
        query = query.filter(Application.job_id == job_id)
        
    # Security: Only admins can see everything. Others see their own jobs' apps.
    if current_user.role != "admin":
        print(f"DEBUG: Filtering applications for HR ID {current_user.id}")
        query = query.filter(Job.hr_id == current_user.id)
    else:
        print("DEBUG: Admin viewing all applications")
        
    applications = query.all()
    print(f"DEBUG: Found {len(applications)} applications for user {current_user.id}")
    
    for app in applications:
        if app.candidate_photo_path and ":" in app.candidate_photo_path:
            idx = app.candidate_photo_path.find("uploads")
            if idx != -1:
                app.candidate_photo_path = app.candidate_photo_path[idx:].replace("\\", "/")
        if app.resume_file_path and ":" in app.resume_file_path:
            idx = app.resume_file_path.find("uploads")
            if idx != -1:
                app.resume_file_path = app.resume_file_path[idx:].replace("\\", "/")
>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
    
    return applications

@router.get("/{application_id}", response_model=ApplicationDetailResponse)
def get_application(
    application_id: int,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get application details (HR only)"""
    application = db.query(Application).options(
        joinedload(Application.job),
        joinedload(Application.resume_extraction),
<<<<<<< HEAD
        joinedload(Application.interview)
=======
        joinedload(Application.interview),
        joinedload(Application.pipeline_stages)
>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
    ).filter(Application.id == application_id).first()
    
    if not application:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Application not found"
        )
    
<<<<<<< HEAD
=======
    # Sanitize paths
    if application.candidate_photo_path and ":" in application.candidate_photo_path:
        idx = application.candidate_photo_path.find("uploads")
        if idx != -1:
            application.candidate_photo_path = application.candidate_photo_path[idx:].replace("\\", "/")
    if application.resume_file_path and ":" in application.resume_file_path:
        idx = application.resume_file_path.find("uploads")
        if idx != -1:
            application.resume_file_path = application.resume_file_path[idx:].replace("\\", "/")
>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
    
    return application

@router.put("/{application_id}/status", response_model=ApplicationDetailResponse)
def update_application_status(
    application_id: int,
    status_update: ApplicationStatusUpdate,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
<<<<<<< HEAD
    """Update application status (HR only)"""
    application = db.query(Application).filter(Application.id == application_id).first()
    
    if not application:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Application not found"
        )
    
    
    # Validate status
    valid_statuses = ["approved_for_interview", "rejected", "review_later"]
    if status_update.status not in valid_statuses:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid status. Must be one of: {valid_statuses}"
        )
    
=======
    """Update application status & Advance Pipeline (Point 1)"""
    application = db.query(Application).filter(Application.id == application_id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    
    from app.services.candidate_service import CandidateService
    cand_service = CandidateService(db)

    # Valid statuses for the simplified HR view (will mapping to pipeline stages)
    valid_statuses = ["approved_for_interview", "rejected", "review_later", "technical_interview", "hr_interview", "hired"]
    if status_update.status not in valid_statuses:
        raise HTTPException(status_code=400, detail=f"Invalid status.")
    
    old_status = application.status
>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
    application.status = status_update.status
    if status_update.hr_notes:
        application.hr_notes = status_update.hr_notes
    
<<<<<<< HEAD
    # Generate Interview Access Key if approved
    raw_access_key = None
    if application.status == "approved_for_interview":
        # Check if interview already exists to prevent duplicates
=======
    # Advance Pipeline Stages (Point 1)
    if status_update.status == "approved_for_interview":
        cand_service.advance_stage(application.id, "Resume Screening", "pass", notes=status_update.hr_notes, evaluator_id=current_user.id)
        cand_service.advance_stage(application.id, "Aptitude Round", "pending")
    elif status_update.status == "technical_interview":
        cand_service.advance_stage(application.id, "Technical Interview", "pending", evaluator_id=current_user.id)
    elif status_update.status == "hr_interview":
        cand_service.advance_stage(application.id, "HR Interview", "pending", evaluator_id=current_user.id)
    elif status_update.status == "hired":
        cand_service.advance_stage(application.id, "Final Decision", "pass", notes=status_update.hr_notes, evaluator_id=current_user.id)
    elif status_update.status == "rejected":
        # Record failure in current stage
        current_stage = "Resume Screening" if old_status == "submitted" else "Final Decision"
        cand_service.advance_stage(application.id, current_stage, "fail", notes=status_update.hr_notes, evaluator_id=current_user.id)

    # Logging
    cand_service.create_audit_log(current_user.id, "STATUS_UPDATED", "Application", application.id, {"from": old_status, "to": status_update.status})

    # Generate Interview Access Key if approved
    raw_access_key = None
    if application.status == "approved_for_interview":
>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
        existing_interview = db.query(Interview).filter(Interview.application_id == application.id).first()
        if not existing_interview:
            import uuid
            raw_access_key = secrets.token_urlsafe(16)
            hashed_key = pwd_context.hash(raw_access_key)
            expiration = datetime.now(timezone.utc) + timedelta(hours=24)
            unique_test_id = f"TEST-{uuid.uuid4().hex[:8].upper()}"
            
            new_interview = Interview(
                test_id=unique_test_id,
                application_id=application.id,
                status='not_started',
                access_key_hash=hashed_key,
                expires_at=expiration,
                is_used=False
            )
            db.add(new_interview)
    
<<<<<<< HEAD
    try:
        db.commit()
        db.refresh(application)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to update application status securely")
    
    # Send email notification to candidate
    candidate_email = application.candidate_email
    job_title = application.job.title
    if application.status == "approved_for_interview" and raw_access_key:
        # Pass raw access key to email
=======
    db.commit()
    db.refresh(application)
    
    # Notifications (Point 9)
    candidate_email = application.candidate_email
    job_title = application.job.title
    if application.status == "approved_for_interview" and raw_access_key:
>>>>>>> fc67732bae97f8da95fde30813676c1c6ceeb92e
        background_tasks.add_task(send_approved_for_interview_email, candidate_email, job_title, raw_access_key)
    elif application.status == "rejected":
        background_tasks.add_task(send_rejected_email, candidate_email, job_title, False)
        
    return application

@router.delete("/{application_id}")
async def delete_application(
    application_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Delete an application along with associated data. HR only.
    """
    if current_user.role != "hr":
        raise HTTPException(status_code=403, detail="Only HR can delete applications")

    app = db.query(Application).filter(Application.id == application_id).first()
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")

    # Explicitly delete related records to avoid constraint violations
    if app.resume_extraction:
        db.delete(app.resume_extraction)
    if app.hiring_decision:
        db.delete(app.hiring_decision)
    if app.interview:
        if app.interview.report:
            db.delete(app.interview.report)
        for question in app.interview.questions:
            db.query(InterviewAnswer).filter(InterviewAnswer.question_id == question.id).delete()
            db.delete(question)
        db.delete(app.interview)

    db.delete(app)
    db.commit()
    return {"message": "Application deleted successfully"}
